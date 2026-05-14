"""Convert sekpriAI_Source_of_Truth_Blueprint.docx to readable Markdown.

Lightweight, single-purpose script. No dependencies beyond python-docx.
This is a Phase 0 utility and not part of the runtime app.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


HEADING_PREFIX = "Heading"


def style_to_md_heading(style_name: str) -> str | None:
    """Map a Word heading style name to a markdown heading prefix."""
    if not style_name:
        return None
    if style_name.startswith(HEADING_PREFIX):
        # e.g. "Heading 1" -> level 1
        try:
            level = int(style_name.replace(HEADING_PREFIX, "").strip() or "1")
            level = max(1, min(level, 6))
            return "#" * level
        except ValueError:
            return "#"
    if style_name == "Title":
        return "#"
    if style_name == "Subtitle":
        return "##"
    return None


def is_list_paragraph(paragraph) -> tuple[bool, bool]:
    """Return (is_list, is_numbered) by inspecting numbering XML."""
    p = paragraph._p
    numpr = p.find(qn("w:pPr"))
    if numpr is None:
        return False, False
    num = numpr.find(qn("w:numPr"))
    if num is None:
        return False, False
    # Heuristic: if style hints "Number" treat as ordered, else bullet.
    style = (paragraph.style.name or "").lower()
    is_numbered = "number" in style or "ordered" in style
    return True, is_numbered


def paragraph_to_md(paragraph) -> str:
    text = paragraph.text.strip()
    if not text:
        return ""

    heading = style_to_md_heading(paragraph.style.name)
    if heading:
        return f"{heading} {text}"

    is_list, is_numbered = is_list_paragraph(paragraph)
    if is_list:
        bullet = "1." if is_numbered else "-"
        return f"{bullet} {text}"

    return text


def table_to_md(table) -> str:
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
    if not rows:
        return ""
    header = rows[0]
    separator = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
    return "\n".join([header, separator, *rows[1:]])


def convert(docx_path: Path, md_path: Path) -> None:
    document = Document(str(docx_path))

    md_lines: list[str] = []
    body = document.element.body

    # Preserve document order: walk through paragraphs and tables in order.
    paragraphs_iter = iter(document.paragraphs)
    tables_iter = iter(document.tables)
    para_map = {p._p: p for p in document.paragraphs}
    table_map = {t._tbl: t for t in document.tables}

    prev_blank = True
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph = para_map.get(child)
            if paragraph is None:
                continue
            line = paragraph_to_md(paragraph)
            if not line:
                if not prev_blank:
                    md_lines.append("")
                    prev_blank = True
                continue
            md_lines.append(line)
            prev_blank = False
        elif child.tag == qn("w:tbl"):
            table = table_map.get(child)
            if table is None:
                continue
            md_lines.append("")
            md_lines.append(table_to_md(table))
            md_lines.append("")
            prev_blank = True

    md_text = "\n".join(md_lines).rstrip() + "\n"
    md_path.parent.mkdir(parents=True, exist_ok=True)
    md_path.write_text(md_text, encoding="utf-8")


def main() -> int:
    src = Path("sekpriAI_Source_of_Truth_Blueprint.docx")
    dst = Path("docs/sekpriAI_Source_of_Truth_Blueprint.md")
    if not src.exists():
        print(f"Source DOCX not found: {src}", file=sys.stderr)
        return 1
    convert(src, dst)
    size_kb = dst.stat().st_size / 1024
    print(f"Converted -> {dst} ({size_kb:.1f} KB)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
